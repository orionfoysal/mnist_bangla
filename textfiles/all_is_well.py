'''
python formCreator.py originFile numberOfFilesToBeCreated
Output folder format: textfiles -> [docs, text]

convert to pdf: 
for f in *.docx; do lowriter --headless --convert-to pdf $f; done
'''


from docx import Document
from docx.shared import Mm
from random import shuffle 
from docx.shared import Pt 
from docx.enum.text import WD_ALIGN_PARAGRAPH
from random import randint
import sys 

fileNames = sys.argv[1:]

text = ''
for ifile in fileNames:
    with open(ifile) as f:
        text += f.read()

all_text = []
lines = text.splitlines()
lines = list(filter(None, lines))
all_text = [line for line in lines if 290 <len(line)< 400]


shuffle(all_text)
# print((all_text))

for para in range(len(all_text)):
    writable = []
    textToWrite = all_text[para]

    with open('textfiles/text/all/'+str(para)+'.txt', 'w') as f:
        f.write(textToWrite)

    
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Lohit Bengali'
    font.size = Pt(14)
    
    section = doc.sections[0]
    section.page_height = Mm(297) 
    section.page_width = Mm(210)
    section.bottom_margin = Mm(12.7)
    section.top_margin = Mm(15)
    section.left_margin = Mm(12.7)
    section.right_margin = Mm(12.7)


    table1 = doc.add_table(rows=4, cols=2)
    table1.style = "TableGrid"
    table1.cell(0,0).text = 'ID'
    table1.cell(1,0).text = 'Occupation'
    table1.cell(2,0).text = 'Gender'
    table1.cell(3,0).text = 'Age'

    table1.cell(0,1).text = '   {:04}'.format(para)
    doc.add_paragraph('\n')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(textToWrite)

    table = doc.add_table(rows=1, cols=1)
    table.style = "TableGrid"
    cell = table.rows[0].cells
    cell[0].text = '\n'*28
    doc.save('textfiles/docs/all/%s.docx'%str(para))
