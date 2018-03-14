'''
python formCreator.py originFile numberOfFilesToBeCreated
Output folder format: textfiles -> [docs, text]

convert to pdf: 
for f in *.docx; do lowriter --headless --convert-to pdf $f; done
'''


from docx import Document
from docx.shared import Mm
from random import shuffle 
from docx.shared import Pt 
from docx.enum.text import WD_ALIGN_PARAGRAPH
from random import randint
import sys 

# fileName = sys.argv[1]
# numberOfFile = int(sys.argv[2])
# with open(fileName) as f:
#     all_text = f.readlines()

all_text = []
numberOfFile = 10
with open('index.txt') as f:
    ind = f.read()
ind = ind.splitlines()
ind = list(filter(None, ind))
# print(ind)
with open('all.txt') as f:
    text = f.read()
lines = text.splitlines()
lines = list(filter(None, lines))
for i in ind:
    for k,line in enumerate(lines):
        if i in line:
            all_text.append(''.join(lines[k+1:k+10]))
            lines = lines[k:]
            break

shuffle(all_text)
# print((all_text))

for para in range(len(all_text)):
    writable = []
    text = ((all_text[para]).split('।'))[0:5]
    for i in text:
        writable.append(i+'।')
    textToWrite = ''.join(writable)
    # textToWrite = all_text[para]    
    # while True:
    #     writable = []
    #     text = (all_text[randint(1,len(all_text))]).split('।')#)#[0:5]
    #     for i in text:
    #         writable.append(i+'।')

        # textToWrite = ''.join(writable)
        # textToWrite += '\"'
        
        # if 33 <len(textToWrite)< 665:
        #     break

    with open('textfiles/text/'+str(para)+'.txt', 'w') as f:
        f.write(textToWrite)

    
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Lohit Bengali'
    font.size = Pt(14)
    
    section = doc.sections[0]
    section.page_height = Mm(297) 
    section.page_width = Mm(210)
    section.bottom_margin = Mm(12.7)
    section.top_margin = Mm(15)
    section.left_margin = Mm(12.7)
    section.right_margin = Mm(12.7)


    table1 = doc.add_table(rows=4, cols=2)
    table1.style = "TableGrid"
    table1.cell(0,0).text = 'ID'
    table1.cell(1,0).text = 'Occupation'
    table1.cell(2,0).text = 'Gender'
    table1.cell(3,0).text = 'Age'

    table1.cell(0,1).text = '   {:04}'.format(para)
    doc.add_paragraph('\n')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(textToWrite)

    table = doc.add_table(rows=1, cols=1)
    table.style = "TableGrid"
    cell = table.rows[0].cells
    cell[0].text = '\n'*25
    doc.save('textfiles/docs/%s.docx'%str(para))
